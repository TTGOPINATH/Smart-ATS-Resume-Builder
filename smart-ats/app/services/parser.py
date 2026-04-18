import io
import re
from typing import Optional


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfminer.six."""
    try:
        from pdfminer.high_level import extract_text
        from pdfminer.layout import LAParams

        pdf_file = io.BytesIO(file_bytes)
        laparams = LAParams(line_margin=0.5, word_margin=0.1)
        text = extract_text(pdf_file, laparams=laparams)
        return text.strip() if text else ""
    except ImportError:
        return "[pdfminer.six not installed — install it to enable PDF parsing]"
    except Exception as e:
        return f"[PDF parsing error: {str(e)}]"


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document

        doc_file = io.BytesIO(file_bytes)
        doc = Document(doc_file)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        return "[python-docx not installed — install it to enable DOCX parsing]"
    except Exception as e:
        return f"[DOCX parsing error: {str(e)}]"


def extract_email(text: str) -> Optional[str]:
    """Extract email address from text."""
    match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    return match.group() if match else None


def extract_phone(text: str) -> Optional[str]:
    """Extract phone number from text."""
    match = re.search(r'[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}', text)
    return match.group() if match else None


def extract_linkedin(text: str) -> Optional[str]:
    """Extract LinkedIn URL from text."""
    match = re.search(r'linkedin\.com/in/[\w\-]+', text, re.I)
    return f"https://{match.group()}" if match else None


def extract_github(text: str) -> Optional[str]:
    """Extract GitHub URL from text."""
    match = re.search(r'github\.com/[\w\-]+', text, re.I)
    return f"https://{match.group()}" if match else None


def extract_name(text: str) -> Optional[str]:
    """Try to extract candidate name (first line heuristic)."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if lines:
        first = lines[0]
        # Name typically 2-4 words, no special chars
        if re.match(r'^[A-Za-z\s\.\-]{3,50}$', first) and len(first.split()) <= 5:
            return first
    return None


def extract_skills_from_text(text: str) -> list:
    """Extract skills section from parsed text."""
    skills = []
    # Look for skills section
    match = re.search(
        r'(?:skills?|technical skills?|core competencies)[:\s]+(.*?)(?:\n\n|\n[A-Z])',
        text, re.I | re.S
    )
    if match:
        skills_text = match.group(1)
        # Split by common separators
        raw = re.split(r'[,|•·\n\t]+', skills_text)
        skills = [s.strip() for s in raw if 2 < len(s.strip()) < 50]
    return skills[:30]


def parse_resume_file(file_bytes: bytes, filename: str) -> dict:
    """
    Main entry: parse file and return structured data.
    """
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        text = parse_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        text = parse_docx(file_bytes)
    elif filename_lower.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="ignore")
    else:
        text = file_bytes.decode("utf-8", errors="ignore")

    parsed = {
        "raw_text": text,
        "full_name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "linkedin": extract_linkedin(text),
        "github": extract_github(text),
        "skills": extract_skills_from_text(text),
    }
    return parsed
